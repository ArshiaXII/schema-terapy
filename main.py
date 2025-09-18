# Required libraries to install via pip:
# pip install google-generativeai langchain-google-genai langchain langchain-chroma chromadb pypdf python-docx python-dotenv fastapi uvicorn python-multipart

import os
import sys
from pathlib import Path
from typing import List, Optional
import logging

# FastAPI imports
from fastapi import FastAPI, HTTPException, status, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import APIKeyHeader
from pydantic import BaseModel

# Environment variables
from dotenv import load_dotenv

# Document processing imports
import PyPDF2
from docx import Document

# LangChain and AI imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_chroma import Chroma


from langchain.chains import RetrievalQA
from langchain.schema import Document as LangChainDocument

# Load environment variables from .env file
load_dotenv()

# Configuration
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")  # Load from environment variable
MY_APP_SECRET_KEY = os.getenv("MY_APP_SECRET_KEY")  # Load API secret key from environment
SOURCE_FOLDER = "kaynaklarim"  # Folder containing PDF and DOCX files
VECTOR_DB_PATH = "chroma_db"  # Path for persistent ChromaDB storage

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# FastAPI app initialization
app = FastAPI(
    title="Schema Therapy RAG API",
    description="""
    A secure FastAPI service for Schema Therapy analysis using RAG (Retrieval-Augmented Generation).

    ## Authentication
    Protected endpoints require an API key to be sent in the `X-API-Key` header.

    ## Endpoints
    - **Public**: `/` (root), `/health` (health check)
    - **Protected**: `/analyze-schemas/`, `/chat-with-results/` (require API key)
    """,
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure this properly for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# API Key Security Configuration
api_key_header = APIKeyHeader(name="X-API-Key", auto_error=False)

async def verify_api_key(api_key: str = Depends(api_key_header)):
    """
    Verify the API key sent in the X-API-Key header.

    Args:
        api_key: The API key from the request header

    Returns:
        The API key if valid

    Raises:
        HTTPException: 401 Unauthorized if the API key is invalid or missing
    """
    if not MY_APP_SECRET_KEY:
        logger.error("MY_APP_SECRET_KEY environment variable not configured!")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Server configuration error: API key not configured"
        )

    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="API key is required. Please provide X-API-Key header.",
            headers={"WWW-Authenticate": "ApiKey"},
        )

    if api_key != MY_APP_SECRET_KEY:
        logger.warning(f"Invalid API key attempt: {api_key[:10]}...")
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid API key",
            headers={"WWW-Authenticate": "ApiKey"},
        )

    return api_key

# Pydantic models for API requests/responses
class SchemaAnalysisRequest(BaseModel):
    schemas: List[str]

class SchemaAnalysisResponse(BaseModel):
    analysis: str
    schemas_analyzed: List[str]

class ChatRequest(BaseModel):
    schemas: List[str]
    question: str

class ChatResponse(BaseModel):
    answer: str
    schemas_context: List[str]
    question: str

# Global variables for the QA chain
qa_chain = None

def setup_environment():
    """Set up the environment with the Google API key and app secret key."""
    if not GOOGLE_API_KEY or GOOGLE_API_KEY.strip() == "":
        logger.error("GOOGLE_API_KEY environment variable not found!")
        logger.error("Please create a .env file with your Google API key.")
        logger.error("You can get one from: https://makersuite.google.com/app/apikey")
        return False

    if not MY_APP_SECRET_KEY or MY_APP_SECRET_KEY.strip() == "":
        logger.error("MY_APP_SECRET_KEY environment variable not found!")
        logger.error("Please add MY_APP_SECRET_KEY to your .env file for API authentication.")
        logger.error("Example: MY_APP_SECRET_KEY=\"your_strong_random_secret_here\"")
        return False

    os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY
    logger.info("Environment variables configured successfully")
    return True

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error reading PDF {pdf_path}: {str(e)}")
        return ""

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error reading DOCX {docx_path}: {str(e)}")
        return ""

def load_documents_from_folder(folder_path: str) -> str:
    """
    Load and extract text from all PDF and DOCX files in the specified folder.
    Returns concatenated text from all documents.
    """
    if not os.path.exists(folder_path):
        logger.error(f"Source folder '{folder_path}' does not exist!")
        logger.error(f"Please create the folder and add your Schema Therapy documents.")
        return ""
    
    all_text = ""
    supported_extensions = ['.pdf', '.docx']
    files_processed = 0
    
    logger.info(f"Loading documents from '{folder_path}'...")
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        file_extension = Path(filename).suffix.lower()
        
        if file_extension in supported_extensions and os.path.isfile(file_path):
            logger.info(f"Processing: {filename}")
            
            if file_extension == '.pdf':
                text = extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                text = extract_text_from_docx(file_path)
            
            if text.strip():
                all_text += f"\n--- Content from {filename} ---\n{text}\n"
                files_processed += 1
            else:
                logger.warning(f"No text extracted from {filename}")
    
    logger.info(f"Successfully processed {files_processed} documents")
    logger.info(f"Total text length: {len(all_text)} characters")
    
    return all_text

def create_vector_database(text_content: str, db_path: str) -> bool:
    """
    Create a ChromaDB vector database from the provided text content.
    Returns True if successful, False otherwise.
    """
    if not text_content.strip():
        logger.error("No text content provided for vector database creation!")
        return False
    
    try:
        logger.info("Creating vector database...")
        
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        logger.info("Splitting text into chunks...")
        chunks = text_splitter.split_text(text_content)
        logger.info(f"Created {len(chunks)} text chunks")
        
        # Create LangChain documents
        documents = [LangChainDocument(page_content=chunk) for chunk in chunks]
        
        # Initialize embeddings
        logger.info("Initializing Google embeddings...")
        embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
        
        # Create and persist vector store
        logger.info("Creating and saving vector database...")
        Chroma.from_documents(
            documents=documents,
            embedding=embeddings,
            persist_directory=db_path
        )
        
        logger.info(f"Vector database created successfully at '{db_path}'")
        return True
        
    except Exception as e:
        logger.error(f"Error creating vector database: {str(e)}")
        return False

def setup_qa_chain(db_path: str) -> Optional[RetrievalQA]:
    """
    Set up the Question-Answering chain with the vector database.
    Returns the QA chain if successful, None otherwise.
    """
    try:
        logger.info("Setting up QA chain...")
        
        # Initialize embeddings
        embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
        
        # Load existing vector store
        logger.info(f"Loading vector database from '{db_path}'...")
        vectorstore = Chroma(
            persist_directory=db_path,
            embedding_function=embeddings
        )
        
        # Initialize the language model
        logger.info("Initializing Gemini model...")
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-pro-latest",
            temperature=0.3,
            convert_system_message_to_human=True
        )
        
        # Create retriever
        retriever = vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 5}
        )
        
        # Create QA chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever,
            return_source_documents=False
        )
        
        logger.info("QA chain setup complete!")
        return qa_chain
        
    except Exception as e:
        logger.error(f"Error setting up QA chain: {str(e)}")
        return None

async def initialize_system():
    """Initialize the RAG system on startup."""
    global qa_chain
    
    logger.info("Initializing Schema Therapy RAG system...")
    
    # Setup environment
    if not setup_environment():
        logger.error("Environment setup failed. Please configure your Google API key.")
        return False
    
    # Check if vector database already exists
    if os.path.exists(VECTOR_DB_PATH):
        logger.info(f"Found existing vector database at '{VECTOR_DB_PATH}'")
        logger.info("Skipping document processing and using existing database...")
    else:
        logger.info(f"Vector database not found at '{VECTOR_DB_PATH}'")
        logger.info("Creating new vector database...")
        
        # Load documents
        combined_text = load_documents_from_folder(SOURCE_FOLDER)
        
        if not combined_text.strip():
            logger.error("No documents found or processed. Please check your source folder.")
            return False
        
        # Create vector database
        if not create_vector_database(combined_text, VECTOR_DB_PATH):
            logger.error("Failed to create vector database.")
            return False
    
    # Setup QA chain
    qa_chain = setup_qa_chain(VECTOR_DB_PATH)
    
    if qa_chain is None:
        logger.error("Failed to setup QA chain.")
        return False
    
    logger.info("Schema Therapy RAG system initialized successfully!")
    return True

# FastAPI Events
@app.on_event("startup")
async def startup_event():
    """Initialize the system when the API starts."""
    success = await initialize_system()
    if not success:
        logger.error("Failed to initialize system. API may not function properly.")

# API Endpoints
@app.get("/")
async def root():
    """Root endpoint with API information."""
    return {
        "message": "Schema Therapy RAG API",
        "version": "1.0.0",
        "description": "A secure FastAPI service for Schema Therapy analysis using RAG",
        "authentication": "Protected endpoints require X-API-Key header",
        "endpoints": {
            "public": {
                "/": "GET - API information (this endpoint)",
                "/health": "GET - Health check endpoint",
                "/docs": "GET - Interactive API documentation"
            },
            "protected": {
                "/analyze-schemas/": "POST - Analyze schemas and generate personalized reports (requires API key)",
                "/chat-with-results/": "POST - Chat about specific schemas and ask follow-up questions (requires API key)"
            }
        },
        "security": {
            "header_name": "X-API-Key",
            "description": "Include your API key in the X-API-Key header for protected endpoints"
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint."""
    global qa_chain
    return {
        "status": "healthy" if qa_chain is not None else "unhealthy",
        "qa_chain_ready": qa_chain is not None,
        "vector_db_exists": os.path.exists(VECTOR_DB_PATH)
    }

@app.post("/analyze-schemas/", response_model=SchemaAnalysisResponse)
async def analyze_schemas(request: SchemaAnalysisRequest, api_key: str = Depends(verify_api_key)):
    """
    Analyze schemas and generate a personalized educational report.

    Args:
        request: SchemaAnalysisRequest containing list of schemas to analyze

    Returns:
        SchemaAnalysisResponse with the generated analysis
    """
    global qa_chain

    if qa_chain is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="QA system not initialized. Please check server logs."
        )

    if not request.schemas:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No schemas provided for analysis."
        )

    try:
        # Create schema list string for the prompt
        schema_list_str = ", ".join(request.schemas)

        # Create the analysis prompt with your specified format
        prompt = f"""
ROLE: You are an expert Schema Therapy assistant. Your tone must be supportive, educational, and non-judgmental. You MUST base your entire analysis strictly on the context provided from the user's uploaded documents. Never use external knowledge. Do not act as a therapist or provide a clinical diagnosis.

CONTEXT: [Your system will automatically place the relevant text snippets from the user's documents here based on the schemas below. You don't need to write anything here.]

USER DATA: A user has completed a questionnaire, and their most dominant schemas have been identified as: **{schema_list_str}**.

TASK: Based ONLY on the provided CONTEXT, generate a personalized and structured educational report for the user. Structure your response using Markdown for clear formatting. For each schema listed in the USER DATA, create a dedicated section. Follow this exact format for each schema:

## ▶️ {{Schema Name}}: General Overview
(Provide a detailed definition and core concept of this schema based on the provided texts.)

### Potential Effects on Your Life
(Based on the texts, explain how this schema might manifest in the user's daily life, thoughts, feelings, and relationships.)

### Next Steps (According to the Sources)
(Summarize the suggestions, strategies, or steps for change related to this schema that are mentioned in the source documents. Always frame this as educational information from the texts, not as direct advice.)

---
(If there is more than one schema, repeat the above structure for the next schema.)

If you cannot find sufficient information for any of the requested schemas in the provided CONTEXT, you must clearly state: "Sağlanan dokümanlarda '{{Schema Name}}' hakkında detaylı bilgi bulunamadı."
"""

        logger.info(f"Analyzing schemas: {schema_list_str}")

        # Get AI response
        result = qa_chain.invoke(prompt)
        analysis_text = result['result']

        logger.info("Schema analysis completed successfully")

        return SchemaAnalysisResponse(
            analysis=analysis_text,
            schemas_analyzed=request.schemas
        )

    except Exception as e:
        logger.error(f"Error during schema analysis: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error during analysis: {str(e)}"
        )

@app.post("/chat-with-results/", response_model=ChatResponse)
async def chat_with_results(request: ChatRequest, api_key: str = Depends(verify_api_key)):
    """
    Chat with the AI about specific schemas and ask follow-up questions.

    Args:
        request: ChatRequest containing user's schemas and follow-up question

    Returns:
        ChatResponse with the AI's contextual answer
    """
    global qa_chain

    if qa_chain is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="QA system not initialized. Please check server logs."
        )

    if not request.schemas:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No schemas provided for context."
        )

    if not request.question.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No question provided."
        )

    try:
        # Create contextual search query combining schemas and question
        schemas_str = ", ".join(request.schemas)
        contextual_query = f"Regarding the schemas '{schemas_str}', the user asks: {request.question}"

        logger.info(f"Processing chat request about schemas: {schemas_str}")
        logger.info(f"User question: {request.question}")

        # Create detailed prompt for contextual chat
        chat_prompt = f"""
You are an expert Schema Therapy assistant providing contextual support and education. Your tone must be supportive, educational, and non-judgmental. You MUST base your entire response strictly on the context provided from the user's uploaded documents. Never use external knowledge.

CONTEXT: [Your system will automatically place the relevant text snippets from the user's documents here based on the contextual query below. You don't need to write anything here.]

USER CONTEXT: The user has been identified with these dominant schemas: **{schemas_str}**.

USER QUESTION: "{request.question}"

CONTEXTUAL QUERY FOR RETRIEVAL: {contextual_query}

TASK: Based ONLY on the provided CONTEXT, provide a helpful, personalized response that:

1. **Directly addresses the user's specific question** in relation to their identified schemas
2. **Provides educational information** from the source documents that relates to both their schemas and their question
3. **Maintains a supportive, non-judgmental tone** throughout the response
4. **Explains concepts clearly** using language that is accessible and understandable
5. **Connects the answer to their specific schema context** when relevant
6. **Avoids providing clinical diagnoses or therapeutic advice** - focus on educational information only
7. **If the source material doesn't contain enough information** to answer the question, clearly state this limitation

Structure your response in a conversational yet informative way. Use the retrieved context to provide accurate, source-based information while keeping the user's specific schemas in mind.

If you cannot find sufficient information in the provided CONTEXT to answer the question, you must clearly state: "Sağlanan dokümanlarda bu sorunuza yeterli bilgi bulunamadı, ancak mevcut bilgiler ışığında şunları söyleyebilirim..."
"""

        # Get AI response using the contextual prompt
        result = qa_chain.invoke(chat_prompt)
        answer_text = result['result']

        logger.info("Chat response generated successfully")

        return ChatResponse(
            answer=answer_text,
            schemas_context=request.schemas,
            question=request.question
        )

    except Exception as e:
        logger.error(f"Error during chat processing: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error during chat processing: {str(e)}"
        )

# Run the application
if __name__ == "__main__":
    print("🚀 Starting API Server...")
    # Import uvicorn locally to ensure it's available.
    import uvicorn

    # Get the port number from the environment variable set by the deployment platform (like Render).
    # Default to 8000 for local development.
    port = int(os.environ.get("PORT", 8000))

    # Run the Uvicorn server.
    # Host '0.0.0.0' is necessary for the service to be accessible from outside its container.
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)
